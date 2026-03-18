#!/usr/bin/env python3
"""
创建汇总文档 - 生成包含汇总表格和原图的Word文档
"""

import json
import argparse
import os
import sys
from pathlib import Path
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import fitz  # PyMuPDF for PDF to image conversion

def setup_logging(verbose=False):
    """设置日志记录"""
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.StreamHandler(sys.stdout)
        ]
    )

def _find_trip_crop_rect(page):
    """检测滴滴行程单中的剪切线，返回剪切线以下有效内容的裁剪区域。
    剪切线是一个横跨页面的扁平图片（宽>70%页宽, 高<20pt）。
    """
    img_info = page.get_image_info()

    scissor_bottom = 0
    for ii in img_info:
        bbox = ii.get('bbox', [0, 0, 0, 0])
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        if w > page.rect.width * 0.7 and 0 < h < 20:
            scissor_bottom = bbox[3]
            break

    if not scissor_bottom:
        return None

    blocks = page.get_text('blocks')
    content_bottom = scissor_bottom
    for b in blocks:
        txt = b[4].strip()
        if b[1] > scissor_bottom and txt and '页码' not in txt:
            content_bottom = max(content_bottom, b[3])

    margin = 15
    return fitz.Rect(
        page.rect.x0,
        scissor_bottom + 2,
        page.rect.x1,
        min(content_bottom + margin, page.rect.height)
    )

def _find_third_party_crop_rect(page):
    """检测首汽约车等第三方网约车行程单内容边界，去除顶部和底部空白。"""
    blocks = page.get_text('blocks')
    content_top = None
    content_bottom = 0

    for b in blocks:
        txt = b[4].strip()
        if not txt:
            continue
        if any(kw in txt for kw in ['首汽约车', '第三方网约车', '行程单', '行程起止日期']):
            if content_top is None:
                content_top = b[1]
        if content_top is not None and '页码' not in txt:
            content_bottom = max(content_bottom, b[3])

    if content_top is None:
        return None

    margin = 15
    return fitz.Rect(
        page.rect.x0,
        max(0, content_top - 5),
        page.rect.x1,
        min(content_bottom + margin, page.rect.height)
    )


def _find_amap_crop_rect(page):
    """检测高德行程单内容边界，去除顶部广告横幅和底部空白。"""
    blocks = page.get_text('blocks')
    content_top = None
    content_bottom = 0

    for b in blocks:
        txt = b[4].strip()
        if not txt:
            continue
        if '高德地图' in txt or 'AMAP' in txt:
            if content_top is None:
                content_top = b[1]
        if content_top is not None and '页码' not in txt:
            content_bottom = max(content_bottom, b[3])

    if content_top is None:
        return None

    margin = 15
    return fitz.Rect(
        page.rect.x0,
        max(0, content_top - 5),
        page.rect.x1,
        min(content_bottom + margin, page.rect.height)
    )


def convert_pdf_to_image(pdf_path, output_dir, dpi=150, crop_type=None):
    """将PDF转换为图片。crop_type='didi'裁剪滴滴广告头, 'amap'裁剪高德横幅和空白。"""
    try:
        doc = fitz.open(pdf_path)

        image_paths = []
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)

            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)

            clip = None
            if crop_type == 'didi':
                clip = _find_trip_crop_rect(page)
            elif crop_type == 'amap':
                clip = _find_amap_crop_rect(page)
            elif crop_type == 'third_party':
                clip = _find_third_party_crop_rect(page)

            pix = page.get_pixmap(matrix=mat, clip=clip) if clip else page.get_pixmap(matrix=mat)

            base_name = Path(pdf_path).stem
            img_filename = f"{base_name}_page_{page_num + 1}.png"
            img_path = Path(output_dir) / img_filename

            pix.save(str(img_path))
            image_paths.append(img_path)

        doc.close()
        return image_paths
    except Exception as e:
        logging.error(f"转换PDF到图片失败 {pdf_path}: {str(e)}")
        return []

def add_header(doc, title="报销凭证汇总"):
    """添加文档标题"""
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(6)
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_summary_table(doc, extractions):
    """创建汇总表格 — 统一表头：序号、条目、时间、起点、终点、金额
    发票如果已配对行程单，则不单独出行（避免金额重复计算）。
    """
    if not extractions:
        logging.info("没有提取到的数据用于创建表格")
        return

    expense_items = [
        ex for ex in extractions
        if not (ex.get('doc_type') == '发票' and ex.get('pair_id'))
    ]

    headers = ["序号", "条目", "时间", "起点", "终点", "金额"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = etree.SubElement(tbl_pr, qn('w:tblW'))
    tbl_w.set(qn('w:type'), 'auto')
    tbl_w.set(qn('w:w'), '0')

    # 自适应列宽：设置 tblLayout=autofit，每个单元格宽度设为 auto
    tbl_layout = tbl_pr.find(qn('w:tblLayout'))
    if tbl_layout is None:
        tbl_layout = etree.SubElement(tbl_pr, qn('w:tblLayout'))
    tbl_layout.set(qn('w:type'), 'autofit')

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        tc_pr = hdr_cells[i]._element.get_or_add_tcPr()
        shd = etree.SubElement(tc_pr, qn('w:shd'))
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'd9ead3')

    total_amount = 0.0
    for idx, ex in enumerate(expense_items, 1):
        row_cells = table.add_row().cells
        etype = ex.get('type', '未知')
        amount = float(ex.get('amount', 0))
        total_amount += amount

        if etype == '12306':
            start = ex.get('departure', '')
            end = ex.get('destination', '')
            label = '12306'
        elif etype == 'amap':
            start = ex.get('start_location', '')
            end = ex.get('end_location', '')
            label = '高德'
        elif etype == 'third_party':
            start = ex.get('start_location', '')
            end = ex.get('end_location', '')
            label = '第三方'
        elif etype in ('dining', '餐饮', '酒店', '机票', '其他发票'):
            start = ex.get('seller_name', '') or ex.get('restaurant_name', '')
            end = ''
            label = ex.get('type', '餐饮') if ex.get('type') != 'dining' else '餐饮'
        else:
            start = ex.get('start_location', '')
            end = ex.get('end_location', '')
            label = '滴滴'

        row_cells[0].text = str(idx)
        row_cells[1].text = label
        row_cells[2].text = ex.get('date', '')
        row_cells[3].text = start
        row_cells[4].text = end
        row_cells[5].text = f"{amount:.2f}"

    total_cells = table.add_row().cells
    total_cells[0].text = ''
    total_cells[1].text = ''
    total_cells[2].text = ''
    total_cells[3].text = ''
    total_cells[4].text = '合计'
    total_cells[5].text = f"{total_amount:.2f}"
    for cell in total_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    # 将所有单元格宽度设为 auto，配合 tblLayout=autofit 实现按内容自适应列宽
    for row in table.rows:
        for cell in row.cells:
            tc_w = cell._element.find(qn('w:tcPr'))
            if tc_w is not None:
                w_el = tc_w.find(qn('w:tcW'))
                if w_el is not None:
                    w_el.set(qn('w:type'), 'auto')
                    w_el.set(qn('w:w'), '0')

def add_images_to_doc(doc, image_paths, max_width=6.0, keep_together=False):
    """向文档添加图片（紧凑布局，配对图片尽量同页）"""
    if not image_paths:
        return

    for i, img_path in enumerate(image_paths):
        if not Path(img_path).exists():
            logging.warning(f"图片文件不存在: {img_path}")
            continue

        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            if keep_together and i < len(image_paths) - 1:
                para.paragraph_format.keep_with_next = True
            run = para.add_run()
            run.add_picture(str(img_path), width=Inches(min(max_width, 6.0)))
        except Exception as e:
            logging.error(f"添加图片失败 {img_path}: {str(e)}")

def add_summary_stats(doc, extractions):
    """添加统计摘要（已配对的发票不重复计算）"""
    if not extractions:
        return

    expense_items = [
        ex for ex in extractions
        if not (ex.get('doc_type') == '发票' and ex.get('pair_id'))
    ]
    total_amount = sum(float(ex.get('amount', 0)) for ex in expense_items)
    total_items = len(expense_items)

    stats_para = doc.add_paragraph()
    stats_para.paragraph_format.space_before = Pt(6)
    stats_para.paragraph_format.space_after = Pt(6)
    stats_run = stats_para.add_run(f"统计摘要: 共 {total_items} 项, 总金额 {total_amount:.2f}元")
    stats_run.font.size = Pt(11)
    stats_run.bold = True

def create_summary_document(info_file, pdf_dir, output_path, config=None):
    """创建汇总文档"""
    if config is None:
        config = {}

    document_config = config.get('document', {})
    processing_config = config.get('processing', {})

    # 默认配置
    default_doc_config = {
        'title': '报销凭证汇总',
        'include_table': True,
        'include_images': True,
        'sort_by': 'date',
        'compact_layout': True
    }
    default_doc_config.update(document_config)

    # 默认处理配置（兼容 processing.image 与 processing.image_conversion）
    default_proc_config = {
        'image': {
            'dpi': 150,
            'width_inches': 6.0,
            'quality': 90
        }
    }
    if 'image' not in default_proc_config:
        default_proc_config['image'] = {}
    default_proc_config['image'].update(processing_config.get('image', {}))
    # test_config 等使用 image_conversion.dpi / max_width_inches
    image_conv = processing_config.get('image_conversion', {})
    if image_conv:
        default_proc_config['image'].setdefault('dpi', image_conv.get('dpi', 150))
        default_proc_config['image'].setdefault('width_inches', image_conv.get('max_width_inches', image_conv.get('width_inches', 6.0)))
    default_proc_config.update(processing_config)

    # 读取信息文件
    with open(info_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    extractions = data.get('extractions', [])

    # 根据日期排序，与 process_expense 一致：(date, pair_id, 行程单优先于发票)
    sort_by = default_doc_config.get('sort_by')
    def _sort_key(x):
        d = x.get('date', '')
        p = x.get('pair_id', '')
        doc = 1 if x.get('doc_type') == '发票' else 0
        return (d, p, doc)
    if sort_by in ('date', 'date_asc'):
        extractions.sort(key=_sort_key)
    elif sort_by == 'date_desc':
        extractions.sort(key=_sort_key, reverse=True)

    # 创建Word文档
    doc = Document()

    # 添加页面边距设置
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # 标题 = 文件名去掉最后的生成日期部分
    stem = Path(output_path).stem
    parts = stem.rsplit('_', 1)
    doc_title = parts[0] if len(parts) > 1 and parts[1].isdigit() and len(parts[1]) == 8 else stem
    add_header(doc, doc_title)

    # 添加汇总表格
    if default_doc_config.get('include_table'):
        create_summary_table(doc, extractions)

    # 添加统计摘要
    add_summary_stats(doc, extractions)

    # 转换PDF为图片并添加到文档（配对的行程单+发票连续粘贴，顺序与表格一致）
    if default_doc_config.get('include_images'):
        import shutil
        temp_img_dir = Path(output_path).parent / "temp_images"
        temp_img_dir.mkdir(exist_ok=True)

        pdf_path = Path(pdf_dir)
        dpi = default_proc_config['image'].get('dpi', 150)
        max_w = default_proc_config['image'].get('width_inches', 6.0)
        added_pairs = set()
        added_files = set()

        def _get_crop_type(ext):
            if ext.get('doc_type') == '行程单':
                t = ext.get('type')
                return t if t in ('didi', 'amap', 'third_party') else None
            return None

        expense_items = [ex for ex in extractions if not (ex.get('doc_type') == '发票' and ex.get('pair_id'))]
        for ex in expense_items:
            pair_id = ex.get('pair_id', '')
            if pair_id:
                if pair_id in added_pairs:
                    continue
                group = [e for e in extractions if e.get('pair_id') == pair_id]
                group.sort(key=lambda e: (1 if e.get('doc_type') == '发票' else 0))
                image_paths = []
                for g in group:
                    g_file = pdf_path / Path(g['filepath']).name
                    if g_file.exists() and g_file.name not in added_files:
                        image_paths.extend(convert_pdf_to_image(
                            g_file, temp_img_dir, dpi=dpi, crop_type=_get_crop_type(g)))
                        added_files.add(g_file.name)
                added_pairs.add(pair_id)
                add_images_to_doc(doc, image_paths, max_w, keep_together=True)
            else:
                pdf_file = pdf_path / Path(ex['filepath']).name
                if not pdf_file.exists():
                    logging.warning(f"未找到PDF文件: {pdf_file}")
                    continue
                if pdf_file.name in added_files:
                    continue
                image_paths = convert_pdf_to_image(
                    pdf_file, temp_img_dir, dpi=dpi, crop_type=_get_crop_type(ex))
                added_files.add(pdf_file.name)
                add_images_to_doc(doc, image_paths, max_w, keep_together=False)

        shutil.rmtree(temp_img_dir)

    # 添加页码（页脚居中: 第 X / Y 页）
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fp.add_run('第 ').font.size = Pt(9)
    r1 = fp.add_run()
    r1.font.size = Pt(9)
    fld1 = etree.SubElement(r1._element, qn('w:fldChar'))
    fld1.set(qn('w:fldCharType'), 'begin')
    r2 = fp.add_run()
    r2.font.size = Pt(9)
    inst1 = etree.SubElement(r2._element, qn('w:instrText'))
    inst1.set(qn('xml:space'), 'preserve')
    inst1.text = ' PAGE '
    r3 = fp.add_run()
    fld2 = etree.SubElement(r3._element, qn('w:fldChar'))
    fld2.set(qn('w:fldCharType'), 'end')

    fp.add_run(' / ').font.size = Pt(9)

    r4 = fp.add_run()
    r4.font.size = Pt(9)
    fld3 = etree.SubElement(r4._element, qn('w:fldChar'))
    fld3.set(qn('w:fldCharType'), 'begin')
    r5 = fp.add_run()
    r5.font.size = Pt(9)
    inst2 = etree.SubElement(r5._element, qn('w:instrText'))
    inst2.set(qn('xml:space'), 'preserve')
    inst2.text = ' NUMPAGES '
    r6 = fp.add_run()
    fld4 = etree.SubElement(r6._element, qn('w:fldChar'))
    fld4.set(qn('w:fldCharType'), 'end')

    fp.add_run(' 页').font.size = Pt(9)

    # 保存文档
    doc.save(output_path)
    logging.info(f"汇总文档已保存: {output_path}")

    return output_path

def main():
    parser = argparse.ArgumentParser(description='创建包含汇总表格和原图的Word文档')
    parser.add_argument('--info', '-i', required=True, help='提取的信息文件路径（JSON格式）')
    parser.add_argument('--pdfs', '-p', required=True, help='包含PDF文件的目录')
    parser.add_argument('--output', '-o', required=True, help='输出Word文档路径')
    parser.add_argument('--config', '-c', help='配置文件路径')
    parser.add_argument('--verbose', '-v', action='store_true', help='显示详细日志')

    args = parser.parse_args()
    setup_logging(args.verbose)

    # 读取配置文件
    config = {}
    if args.config:
        try:
            with open(args.config, 'r', encoding='utf-8') as f:
                config = json.load(f)
        except FileNotFoundError:
            logging.error(f"配置文件不存在: {args.config}")
        except json.JSONDecodeError:
            logging.error(f"配置文件格式错误: {args.config}")

    try:
        logging.info(f"开始创建汇总文档...")
        logging.info(f"信息文件: {args.info}")
        logging.info(f"PDF目录: {args.pdfs}")
        logging.info(f"输出文件: {args.output}")

        create_summary_document(args.info, args.pdfs, args.output, config)

        logging.info("汇总文档创建完成")

    except Exception as e:
        logging.error(f"程序执行出错: {str(e)}")
        sys.exit(1)

if __name__ == '__main__':
    main()