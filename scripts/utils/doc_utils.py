"""
Word文档处理工具函数
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import logging


def create_document_with_settings():
    """创建带有预设格式的Word文档"""
    doc = Document()

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    return doc


def add_title(doc, title, size=Pt(16)):
    """添加文档标题"""
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = size
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加空行
    doc.add_paragraph()

    return title_para


def create_table_with_headers(doc, headers, data_rows=None):
    """创建带标题的表格"""
    if not headers:
        return None

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        # 设置表头背景色
        shading_elm = qn('w:shd')
        cell_properties = hdr_cells[i]._element.get_or_add_tcPr()
        cell_properties.set(shading_elm, 'auto', 'd9ead3')

    # 填充数据行
    if data_rows:
        for row_data in data_rows:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row_data):
                if i < len(row_cells):  # 确保不会超出表格列数
                    row_cells[i].text = str(cell_value)

    return table


def add_centered_paragraph(doc, text, size=None, bold=False):
    """添加居中段落"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)

    if size:
        run.font.size = size
    if bold:
        run.bold = True

    return para


def add_image_to_doc(doc, image_path, max_width=Inches(6.0)):
    """向文档添加图片"""
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()

        # 添加图片，限制最大宽度
        run.add_picture(str(image_path), width=max_width)

        return para
    except Exception as e:
        logging.error(f"添加图片失败 {image_path}: {str(e)}")
        return None


def apply_cell_styling(cell, text_alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """应用单元格样式"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = text_alignment


def add_footer(doc, footer_text):
    """添加页脚"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = footer_text
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_break(doc):
    """添加分节符"""
    doc.add_section()


def adjust_table_column_widths(table, widths):
    """调整表格列宽（需要更底层的操作）"""
    # 这个功能需要使用docx的底层API，比较复杂
    # 当前简化处理，实际应用中可能需要更详细的列宽控制
    pass


def save_document(doc, filepath):
    """保存文档"""
    try:
        doc.save(filepath)
        logging.info(f"文档已保存: {filepath}")
        return True
    except Exception as e:
        logging.error(f"保存文档失败 {filepath}: {str(e)}")
        return False


def add_formatted_text_to_cell(cell, text, bold=False, size=None):
    """向单元格添加格式化文本"""
    paragraphs = cell.paragraphs
    if not paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = paragraphs[0]

    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = size

    return run